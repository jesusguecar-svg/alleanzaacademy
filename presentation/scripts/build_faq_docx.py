#!/usr/bin/env python3
"""Genera el FAQ para Directores en Word (.docx) con branding Alleanza Academy."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ASSETS = os.path.join(ROOT, "assets")
REPO = os.path.dirname(ROOT)

PURPLE = RGBColor(0x7C, 0x46, 0xA6)
NAVY   = RGBColor(0x06, 0x13, 0x30)
GREY   = RGBColor(0x5B, 0x60, 0x70)
FONT   = "Calibri"   # fallback universal; Inter si está instalada

doc = Document()
st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(10.5)
st.font.color.rgb = NAVY

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor); tcPr.append(sh)

def para(text="", size=10.5, color=NAVY, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align: p.alignment = align
    if text:
        r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = color
        r.bold = bold; r.italic = italic; r.font.name = FONT
    return p

def h1(text):
    para(text, size=15, color=PURPLE, bold=True, before=12, after=4)
def h2(text):
    para(text, size=12, color=NAVY, bold=True, before=8, after=2)
def q(qtext, atext):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    r = p.add_run(qtext); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = NAVY
    r2 = p.add_run("  " + atext); r2.font.size = Pt(10.5); r2.font.color.rgb = GREY

# ---- Encabezado con logo ----
logo = os.path.join(ASSETS, "alleanza_logo.png")
if os.path.exists(logo):
    doc.add_picture(logo, width=Inches(2.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
para("FAQ para Directores", size=22, color=NAVY, bold=True, after=2, before=6)
para("Cómo presentar, comunicar y manejar Alleanza Academy ante los agentes",
     size=11.5, color=PURPLE, bold=True, after=10)

# ---- Mensaje en una frase ----
h1("El mensaje en una frase")
para('"Alleanza Academy es nuestra academia interna para que cada agente obtenga su licencia '
     'de Salud y Vida, preparado de verdad para aprobar el examen — en español, a su ritmo y '
     'sin salir de la organización."', italic=True, color=NAVY)
para("Tres ideas que nunca deben faltar:", bold=True, after=2)
for t in ["Propósito: que el agente apruebe el examen (no “tomar un curso bonito”).",
          "In-house: todo dentro de Alleanza; el agente no busca afuera.",
          "En español, con los términos clave en inglés."]:
    doc.add_paragraph(t, style="List Bullet")

# ---- A ----
h1("A. Sobre la academia")
q("¿Qué es Alleanza Academy?",
  "La academia interna de la organización. Su primer producto es el curso de preparación para el examen de licencia de seguros de Salud y Vida.")
q("¿Por qué la creamos si ya existen Kaplan, Xcel, etc.?",
  "Porque esos cursos están solo en inglés y casi todos nuestros agentes hablan español; y porque están hechos para cumplir un requisito, no para que el agente apruebe. Nosotros nos enfocamos en aprobar, en español y dentro de casa.")
q("¿Es un negocio para vender cursos?",
  "No es el objetivo. La meta es resolver una necesidad real del agente (licenciarse y producir antes). El ingreso grande viene de la producción, no de los cursos.")

# ---- B ----
h1("B. El curso")
para("Incluye:", bold=True, after=2)
for t in ["Simulador de examen real — el núcleo del método.",
          "Clases en vivo por Zoom.",
          "Lecciones en video interactivo.",
          "Flashcards de repaso.",
          "Audio / podcasts.",
          "Material de lectura estructurado."]:
    doc.add_paragraph(t, style="List Bullet")
q("¿Por qué tanto énfasis en el simulador?",
  "Porque aprobar depende de practicar en condiciones reales. Todo lo demás refuerza ese objetivo.")
q("¿En qué idioma es?", "En español, con los términos clave en inglés.")

# ---- C ----
h1("C. Rutas, tiempos y acceso")
q("¿Cuánto dura?",
  "Dos planes: ruta intensiva (≈2 semanas al examen) y ruta de un mes (≈1 mes, con más profundidad).")
q("¿Cuánto acceso a la plataforma?",
  "1 mes para estudiar + 2 semanas para practicar y presentar el examen.")
q("¿Y si necesita más tiempo?", "Puede extender por $50 cada 2 semanas (primera extensión).")

# ---- D ----
h1("D. Precio y promoción")
q("¿Cuánto cuesta?",
  "Valor de mercado ~$350. Por lanzamiento y apoyo a las agencias de Alleanza, un código promocional lo deja en $107.")
q("¿Quién puede usar la promoción?", "Exclusiva para agentes de Alleanza.")
q("¿Cómo se aplica?",
  "El agente entra al curso en TutorLMS, aplica el código y paga en la pasarela el precio con descuento.")
para("Cómo comunicar el precio: primero el propósito y el valor; el precio se presenta como un "
     "beneficio, no como un gasto. Evita liderar con el número.", italic=True, color=PURPLE, before=2)

# ---- E (tabla) ----
h1("E. Servicios opcionales (a-la-carte)")
para("No están incluidos en el precio del curso. Quien los necesite nos contacta directamente. "
     "Se cobran aparte para no inflar el precio de entrada.", after=4)
rows = [("Servicio","Costo","Nota"),
        ("Curso de pre-licencia (Xcel)","$30","Requisito en Florida (en Texas no aplica). Generalmente en inglés."),
        ("Cita de examen (Pearson VUE)","$50","Agendamiento del examen oficial."),
        ("Cita de huellas","Según proveedor","Costo del proveedor."),
        ("Asistencia administrativa","$50","Acompañamiento para coordinar el proceso.")]
tbl = doc.add_table(rows=len(rows), cols=3); tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl.style = "Table Grid"
for ri, row in enumerate(rows):
    for ci, val in enumerate(row):
        cell = tbl.cell(ri, ci); cell.text = ""
        p = cell.paragraphs[0]; r = p.add_run(val)
        r.font.size = Pt(9.5); r.font.name = FONT
        if ri == 0:
            r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); shade(cell, "7C46A6")
        else:
            r.font.color.rgb = NAVY if ci < 2 else GREY
            if ci == 1: r.bold = True

# ---- F ----
h1("F. Pre-licencia (manejar con honestidad)")
q("¿El curso cumple el requisito de pre-licencia?",
  "No. A diferencia de Kaplan/Xcel, no es el curso oficial de pre-licencia. Nuestro foco es prepararte para aprobar el examen.")
q("¿Y si el agente necesita la pre-licencia (Florida)?",
  "Lo apoyamos con el curso de pre-licencia de Xcel como servicio opcional ($30). En Texas no se requiere.")
para("Regla de oro: nunca prometas que el curso “cumple el requisito de pre-licencia”. "
     "Sé claro: preparamos para aprobar; la pre-licencia es un apoyo aparte cuando aplica.",
     italic=True, color=PURPLE, before=2)

# ---- G (tabla) ----
h1("G. Ventaja competitiva")
crows = [("","Alleanza Academy","Kaplan / Xcel","María Papi / Aaron"),
         ("Propósito","Que apruebes","Cumplir requisito","Dar clase / repaso"),
         ("Idioma","Español (términos en inglés)","Solo inglés","Español"),
         ("Simulador realista","Sí — núcleo","Genérico","Pocas diapositivas / no"),
         ("Formatos","Zoom + video + flashcards + podcast + lectura","Online autoguiado","Un solo formato"),
         ("Inversión","$107 (valor ~$350)","Mayor","~$100 / mensual")]
ct = doc.add_table(rows=len(crows), cols=4); ct.style = "Table Grid"
for ri, row in enumerate(crows):
    for ci, val in enumerate(row):
        cell = ct.cell(ri, ci); cell.text = ""
        p = cell.paragraphs[0]; r = p.add_run(val); r.font.size = Pt(9); r.font.name = FONT
        if ri == 0:
            r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); shade(cell, "06133 0".replace(" ",""))
        elif ci == 1:
            r.bold = True; r.font.color.rgb = PURPLE; shade(cell, "EAECF3")
        elif ci == 0:
            r.bold = True; r.font.color.rgb = NAVY
        else:
            r.font.color.rgb = GREY
para("Cómo hablar de la competencia: con respeto. No descalificamos a nadie; resaltamos "
     "nuestro enfoque (aprobar, en español, simulador real, dentro de Alleanza).",
     italic=True, color=GREY, before=4)

# ---- H ----
h1("H. Inscripción y soporte")
q("¿Cómo se inscribe un agente?",
  "En TutorLMS → aplica el código promocional → paga en la pasarela. Tras pagar, accede a la plataforma y materiales.")
q("¿Quién da el acompañamiento?",
  "El instructor de la academia (actualmente Jesús Esteban Guerra), con seguimiento directo durante la ruta.")
q("¿A quién escalo casos especiales (huellas, examen, pre-licencia)?",
  "Directo con el equipo de la academia, que coordina los servicios opcionales.")

# ---- I ----
h1("I. Qué decir / qué NO prometer")
h2("Sí comunicamos")
for t in ['"Te preparamos para aprobar, en español."',
          '"Simulador real, clases en vivo y materiales a tu ritmo."',
          '"Rutas de 2 semanas o 1 mes según tu urgencia."',
          '"Precio de lanzamiento de $107, exclusivo para agentes de Alleanza."']:
    doc.add_paragraph(t, style="List Bullet")
h2("No prometemos")
for t in ['Que "cumple el requisito de pre-licencia" (no lo cumple; es apoyo aparte).',
          "Aprobación garantizada (preparamos para aprobar; el esfuerzo del agente cuenta).",
          "Que los servicios a-la-carte están incluidos en el precio del curso."]:
    doc.add_paragraph(t, style="List Bullet")

# ---- J ----
h1("J. Preguntas rápidas de agentes")
qa = [("¿Está en español?","Sí, en español con los términos clave en inglés."),
      ("¿En cuánto tiempo me preparo?","En 2 semanas (intensivo) o 1 mes, tú eliges."),
      ("¿Cuánto cuesta?","$107 con el código de agentes de Alleanza (valor ~$350)."),
      ("¿Me ayudan con el examen y las huellas?","Sí, como apoyo aparte; contáctanos."),
      ("¿Esto es el curso de pre-licencia?","No; preparamos para aprobar. La pre-licencia (FL) la apoyamos por separado."),
      ("¿Y si no me alcanza el tiempo?","Puedes extender el acceso por $50 cada 2 semanas.")]
for qq, aa in qa: q(qq, aa)

para("Recordatorio final: todos comunicamos lo mismo — propósito (aprobar), en español, dentro "
     "de Alleanza. Ante cualquier duda que no esté aquí, escalar al equipo de la academia antes "
     "de improvisar.", italic=True, color=PURPLE, before=12)

out = os.path.join(REPO, "FAQ_Directores_Alleanza_Academy.docx")
doc.save(out)
print("saved", out)
